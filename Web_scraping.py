from bs4 import BeautifulSoup
import requests
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
from docx import Document
from docx.shared import Pt
import datetime

def get_data(link):
    r = requests.get(link,stream=True,verify = False)#, proxies=proxies)
    content = r.content
    soup = BeautifulSoup(content,'html.parser')
    return soup



S = get_data("http://www.mapexpress.ma/?s=bourita")
res_final = []
df = datetime.datetime(2021,10,7)
dict_date = {"janvier":1,"février":2,"mars":3,"avril":4,"mai":5,"juin":6,"juillet":7,"août":8,"septembre":9,"octobre":10,"novembre":11,"décembre":12}
def f(S):
    res = []
    for x in S.findAll("div",attrs = {"class":"single"}):
        d_c = x.find("div",attrs = {"class":"pull-left"}).text[9:]
        jour = d_c[0:2].strip()
        Année = d_c[-4:].strip()
        mois = d_c[2:len(d_c)-4].strip()
        d =  datetime.datetime(int(Année),dict_date[mois],int(jour))
        if df<=d:
            res.append(x.find("a").get("href"))
    return res

def liste_lien_ministre(Ministre_nom,n):
    res_final = []
    for x in range(1,n):
        S = get_data("http://www.mapexpress.ma/page/"+str(x)+"/?s="+Ministre_nom)
        res_final = res_final + f(S)
    return res_final
    

liste_des_ministres = ["AKHANNOUCH","Bourita","LAFTIT","OUAHBI","TOUFIQ","HAJOUI","FETTAH","BARAKA","BENMOUSSA","AIT TALEB","EL MANSOURI","SADIKI","SEKKOURI","MEZZOUR","Fatima Zahra AMMOR","MIRAOUI","BENALI","ABDELJALIL","BENSAID","HAYAR","LOUDYI","JAZOULI","LEKJAA","BAITAS","Ghita MEZZOUR"]



S = get_data("http://www.mapexpress.ma/actualite/activite-gouvernementale/majorite-gouvernementale-homogene-compacte-m-akhannouch/")
#1-Le titre
def titre(S):
    return S.find('div',attrs = {'class':"body single_act"}).find("h1").text

date_dic = {"janvier":"01","février":"02","mars":"03","avril":"04","mai":"05","juin":"06","juillet":"07","août":"08","septembre":"09","octobre":"10","novembre":"11","décembre":"12"}
def date_heure(S):
    text_date = S.find('div',attrs = {'class':'a-post'}).text
    m = text_date.split()
    jour = m[1]
    mois = date_dic[m[2][:-1]]
    Année = m[3]
    date_de_publication = jour+'/'+mois+'/'+Année
    heure_de_publication = m[5]
    return date_de_publication,heure_de_publication

#verifie si l'image existe ou pas 
def check_image(S):
    r = S.find('div',attrs = {'class':"body single_act"})
    if (r is not None):
        response = requests.get(r.find_all("img")[0].get("src"),verify=False,stream=True)
        file = open("sample_i.png", "wb")
        file.write(response.content)
        file.close()
        return True
    else: 
        return False

check_image(S)

#4-Liste des paragraphes gras et italic 
def Paragraphes(S):
    i = 0 
    d = {} #dictionnaire des paragraphes indexés 
    for x in S.find("div",attrs={'class':'body single_act'}).findAll("p"):
        d[i]=x.text
        i = i+1
    return d

def ecrire_article(S,doc):
    #Titre
    H = doc.add_heading(titre(S), 1)
    title_style = H.style
    title_style.font.size = Pt(24)
    rFonts = title_style.element.rPr.rFonts
    rFonts.set(qn("w:asciiTheme"), "Garamond")
    H.alignment = 3     

    #Date et heure de publication
    Date_p,Heure_p = date_heure(S)
    p = doc.add_paragraph('')
    p.add_run('Date de publication ').bold = True
    p.add_run(": "+Date_p)
    p.alignment = 3
    p = doc.add_paragraph('')
    p.add_run('Heure de publication ').bold = True
    p.add_run(": "+Heure_p)
    p.alignment = 3

    #Image
    if(check_image(S)):
        doc.add_picture('sample_i.png', width=Inches(5.9488189), height=Inches(3.594488))


    #Remplissage des paragraphes
    Dict_parag = Paragraphes(S)
    for x in Dict_parag:
        p = doc.add_paragraph('')
        p.add_run(Dict_parag[x]) 
        p.alignment = 3 # for left, 1 for center, 2 right, 3 justify ....
    doc.add_page_break()



def Ministre_RDP(Nom_ministre,n):
    res = liste_lien_ministre(Nom_ministre,n)
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Garamond'
    font.size = Pt(18)
    for x in res:
        S = get_data(x)
        ecrire_article(S,doc)
    doc.save('C:/Users/HP/Desktop/KBscan/RDP/RDP API/Algerie_press_service/'+Nom_ministre+'.docx')

dict_ministre_nb_article = {"BOURITA":16,
                            "LAFTIT": 16,
                            "":16,
                            "":16,
                            "":16,
                            "":16,}
Ministre_RDP(liste_des_ministres[1])
Ministre_RDP(liste_des_ministres[2])
Ministre_RDP(liste_des_ministres[3])
Ministre_RDP(liste_des_ministres[4])
Ministre_RDP(liste_des_ministres[5])
Ministre_RDP(liste_des_ministres[6])
Ministre_RDP(liste_des_ministres[7],5)
Ministre_RDP(liste_des_ministres[8],16)
Ministre_RDP(liste_des_ministres[9],16)
Ministre_RDP(liste_des_ministres[10],4)
Ministre_RDP(liste_des_ministres[11],16)
Ministre_RDP(liste_des_ministres[12],4)
Ministre_RDP(liste_des_ministres[13],10)
Ministre_RDP(liste_des_ministres[14],4)
Ministre_RDP(liste_des_ministres[15],6)

Ministre_RDP(liste_des_ministres[16],4)
Ministre_RDP(liste_des_ministres[17],4)
Ministre_RDP(liste_des_ministres[18],5)
Ministre_RDP(liste_des_ministres[19],6)

Ministre_RDP(liste_des_ministres[20],3)

Ministre_RDP(liste_des_ministres[22],3)

Ministre_RDP(liste_des_ministres[23],19)

Ministre_RDP(liste_des_ministres[21],3)
dict_ministre = {"M. Nasser BOURITA":"Ministre des Affaires étrangères, de la coopération africaine et des Marocains résidant à l’étranger",
                 "M. Aziz AKHANNOUCH":"Chef du gouvernement",
                 "M. Abdelouafi LAFTIT":"Ministre de l'Intérieur",
                 "M. Abdellatif OUAHBI":"Ministre de la Justice",
                 "M. Ahmed TOUFIQ":"Ministre des Habous et des Affaires islamiques",
                 "M. Mohamed HAJOUI":"Secrétaire général de gouvernement",
                 "Mme Nadia FETTAH":"Ministre de l’Economie et des Finances",
                 "M. Chakib BENMOUSSA":"Ministre de l’Education nationale, du Préscolaire et des Sports",
                 "M. Khalid AIT TALEB":"Ministre de la Santé et de la Protection sociale",
                 "M. Mohammed SADIKI":"Ministre de l’Agriculture, de la Pêche maritime/, du Développement rural et des Eaux et Forêts",
                 "M. Younes SEKKOURI OUBBAHESSOU":"Ministre de l’Inclusion économique/, de la Petite entreprise, de l’Emploi et des Compétences",
                 "Mme Fatima Zahra AMMOR":"Ministre du Tourisme, de l’Artisanat et de l’Economie sociale et solidaire",
                 "M. Abdellatif MIRAOUI":"Ministre de l’Enseignement supérieur, de la recherche scientifique et de l’innovation",
                 "Mme Leila BENALI":"Ministre de la Transition énergétique et du Développement durable",
                 "M. Mohammed ABDELJALIL":"Ministre du Transport et de la Logistique",
                 "M. Mohammed Mehdi BENSAID":"Ministre de la Jeunesse, de la Culture et de la Communication",
                 "Mme Aawatif HAYAR":"Ministre de la Solidarité, de l’Insertion sociale et de la Famille",
                 "M. Abdeltif LOUDYI":"Ministre délégué auprès du Chef du gouvernement chargé de l’administration de la Défense nationale",
                 "M. Mohcine JAZOULI":"Ministre délégué auprès du chef du gouvernement chargé de l’Investissement, de la Convergence et de l’évaluation des Politiques publiques",
                 "M. Fouzi LEKJAA":"Ministre délégué auprès du ministre de l’Economie et des Finances, chargé du Budget",
                 "M. Mustapha BAITAS":"Ministre délégué auprès du chef du gouvernement chargé des Relations avec le parlement/, Porte-parole du gouvernement",
                 "Mme Fatim Ezzahra EL MANSOURI":"Ministre de l’Aménagement du territoire national, de l’Urbanisme, de l’Habitat et de la Politique de la ville"}


dict_ministre_nb_article = {"BOURITA":16,
                            "LAFTIT": 16,
                            "OUAHBI":16,
                            "TOUFIQ":16,
                            "HAJOUI":16,
                            "AKHANNOUCH":16,
                            "FETTAH":16,
                            "BENMOUSSA":16,
                            "TALEB":16,
                            "MANSOURI":4,
                            "SADIKI":16,
                            "OUBBAHESSOU":4,
                            "AMMOR":4,
                            "BENALI":4,
                            "MIRAOUI":6,
                            "ABDELJALIL":1,
                            "BENSAID":5,
                            "HAYAR":6,
                            "LOUDYI":1,
                            "JAZOULI":3,
                            "LEKJAA":1,
                            "BAITAS":19,
                            }



